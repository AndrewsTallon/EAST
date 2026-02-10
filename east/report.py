"""Professional Word document report generator for EAST tool."""

import io
import os
import logging
from datetime import datetime
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from east.config import EASTConfig
from east.tests.base import TestResult
from east.visuals.badges import COLORS, create_grade_badge, create_score_gauge, create_status_badge
from east.visuals.tables import (
    create_professional_table,
    create_status_table,
    set_cell_shading,
)

logger = logging.getLogger(__name__)

# Document styling constants
FONT_HEADING = "Calibri"
FONT_BODY = "Calibri"
COLOR_HEADER = RGBColor(0x2C, 0x3E, 0x50)
COLOR_SUBHEADER = RGBColor(0x34, 0x49, 0x5E)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x6C, 0x75, 0x7D)
COLOR_SUCCESS = RGBColor(0x28, 0xA7, 0x45)
COLOR_WARNING = RGBColor(0xFF, 0xC1, 0x07)
COLOR_CRITICAL = RGBColor(0xDC, 0x35, 0x45)
COLOR_INFO = RGBColor(0x17, 0xA2, 0xB8)

SEVERITY_COLORS = {
    "critical": COLOR_CRITICAL,
    "warning": COLOR_WARNING,
    "info": COLOR_INFO,
    "success": COLOR_SUCCESS,
}

SEVERITY_LABELS = {
    "critical": "CRITICAL",
    "warning": "WARNING",
    "info": "INFO",
    "success": "PASS",
}


class EASTReportGenerator:
    """Generates professional Word document reports for EAST scans."""

    OVERVIEW_COLUMNS = 4
    OVERVIEW_ICON_HEIGHT = Inches(0.55)
    OVERVIEW_ROW_HEIGHT = Inches(0.7)
    OVERVIEW_LABEL_FONT_SIZE = Pt(11)

    def __init__(self, config: EASTConfig):
        self.config = config
        self.document = Document()
        self.results: dict[str, list[TestResult]] = {}
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles for professional appearance."""
        style = self.document.styles["Normal"]
        font = style.font
        font.name = FONT_BODY
        font.size = Pt(10)
        font.color.rgb = COLOR_BODY

        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.space_before = Pt(0)

        # Heading styles
        for level in range(1, 4):
            style_name = f"Heading {level}"
            if style_name in self.document.styles:
                h_style = self.document.styles[style_name]
                h_font = h_style.font
                h_font.name = FONT_HEADING
                h_font.color.rgb = COLOR_HEADER
                h_font.bold = True

                if level == 1:
                    h_font.size = Pt(22)
                    h_style.paragraph_format.space_before = Pt(18)
                    h_style.paragraph_format.space_after = Pt(10)
                elif level == 2:
                    h_font.size = Pt(16)
                    h_style.paragraph_format.space_before = Pt(14)
                    h_style.paragraph_format.space_after = Pt(8)
                else:
                    h_font.size = Pt(13)
                    h_style.paragraph_format.space_before = Pt(10)
                    h_style.paragraph_format.space_after = Pt(6)

        # Set default section margins
        for section in self.document.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def add_results(self, domain: str, results: list[TestResult]):
        """Add test results for a domain."""
        self.results[domain] = results

    def generate(self, output_path: str):
        """Generate the complete report and save to file."""
        logger.info("Generating report: %s", output_path)

        self.create_cover_page()
        self._add_page_break()
        self.add_executive_summary()
        self._add_page_break()
        self._add_toc_placeholder()
        self._add_page_break()

        # Add results for each domain
        for domain, results in self.results.items():
            self._add_domain_section(domain, results)

        self._add_appendix()

        # Save
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        self.document.save(output_path)
        logger.info("Report saved to: %s", output_path)

    def create_cover_page(self):
        """Create the cover page with branding."""
        # Add significant spacing at top
        for _ in range(6):
            p = self.document.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Add logo if exists
        logo_path = self.config.branding.logo
        if logo_path:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            self._safe_add_picture(run, logo_path, width=Inches(2.0), warning_context="cover logo")

        # Title
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        run = p.add_run("External Attack Surface Test")
        run.font.size = Pt(32)
        run.font.color.rgb = COLOR_HEADER
        run.font.name = FONT_HEADING
        run.bold = True

        # Subtitle
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("(EAST)")
        run.font.size = Pt(24)
        run.font.color.rgb = COLOR_SUBHEADER
        run.font.name = FONT_HEADING

        # Subtitle line
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Security Assessment Report")
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_MUTED
        run.font.name = FONT_HEADING

        # Horizontal line
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 50)
        run.font.color.rgb = COLOR_INFO
        run.font.size = Pt(10)

        # Spacing
        self.document.add_paragraph()

        # Client info
        client_name = self.config.client_info.name
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Prepared for")
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_MUTED
        run.font.name = FONT_HEADING

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(client_name)
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_HEADER
        run.font.name = FONT_HEADING
        run.bold = True

        # Date
        self.document.add_paragraph()
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime("%B %d, %Y"))
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_MUTED
        run.font.name = FONT_HEADING

        # Company name
        self.document.add_paragraph()
        self.document.add_paragraph()
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.config.branding.company_name)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_MUTED
        run.font.name = FONT_HEADING

        # Confidential notice
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONFIDENTIAL")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_CRITICAL
        run.font.name = FONT_HEADING
        run.bold = True

    def add_executive_summary(self):
        """Add executive summary section."""
        self.document.add_heading("Executive Summary", level=1)

        # Introduction paragraph
        domains_list = ", ".join(self.config.domains) if self.config.domains else "N/A"
        p = self.document.add_paragraph()
        run = p.add_run(
            f"This report presents the findings of an External Attack Surface Test (EAST) "
            f"conducted for {self.config.client_info.name}. The assessment evaluated the "
            f"security posture of the following domains: {domains_list}."
        )
        run.font.size = Pt(10)

        p = self.document.add_paragraph()
        run = p.add_run(
            "The EAST assessment covers SSL/TLS configuration, security headers, "
            "and other externally visible security indicators. The goal is to identify "
            "potential vulnerabilities and provide actionable recommendations."
        )
        run.font.size = Pt(10)

        # Summary dashboard
        if self.results:
            self._add_summary_dashboard()

        # Findings summary
        self._add_findings_summary()

    def _add_summary_dashboard(self):
        """Add a DOCX-native assessment overview grid for each domain."""
        self.document.add_heading("Assessment Overview", level=2)

        for domain, results in self.results.items():
            overview_items = self._build_overview_items(results)
            if not overview_items:
                continue

            p = self.document.add_paragraph()
            run = p.add_run(f"Domain: {domain}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_HEADER
            run.font.name = FONT_HEADING

            self._add_overview_table(overview_items)

    def _build_overview_items(self, results: list[TestResult]) -> list[dict[str, Any]]:
        """Build normalized overview entries for a domain."""
        items: list[dict[str, Any]] = []
        for result in results:
            if result.success and result.score is not None:
                items.append({
                    "label": self._format_test_label(result.test_name),
                    "badge": create_grade_badge(result.grade or "N/A", size=0.6),
                })
            elif not result.success:
                items.append({
                    "label": f"{self._format_test_label(result.test_name)} (ERROR)",
                    "badge": create_status_badge("ERROR", status="critical", size=(1.2, 0.45)),
                })

        return items

    def _add_overview_table(self, overview_items: list[dict[str, Any]]):
        """Render assessment overview icons and labels in an inline table layout."""
        columns = self.OVERVIEW_COLUMNS
        rows = (len(overview_items) + columns - 1) // columns
        table = self.document.add_table(rows=rows, cols=columns)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for row in table.rows:
            row.height = self.OVERVIEW_ROW_HEIGHT
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for index, item in enumerate(overview_items):
            row_idx = index // columns
            col_idx = index % columns
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.width = Inches(1.45)

            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15

            badge_run = paragraph.add_run()
            added = self._safe_add_picture(
                badge_run,
                item["badge"],
                height=self.OVERVIEW_ICON_HEIGHT,
                warning_context=f"overview badge for {item['label']}",
            )

            if added:
                paragraph.add_run("\u00A0\u00A0\u00A0")

            label_run = paragraph.add_run(item['label'])
            label_run.font.name = FONT_BODY
            label_run.font.size = self.OVERVIEW_LABEL_FONT_SIZE
            label_run.font.color.rgb = COLOR_BODY

        self.document.add_paragraph()

    @staticmethod
    def _format_test_label(test_name: str) -> str:
        """Map internal test names to concise overview labels."""
        labels = {
            "ssl_labs": "SSL/TLS",
            "mozilla_observatory": "Observatory",
            "dns_lookup": "DNS",
            "email_auth": "Email Auth",
            "blacklist": "Blacklist",
            "subdomains": "Subdomains",
            "security_headers": "Security Headers",
            "performance": "Performance",
            "cookies": "Cookies",
            "open_ports": "Open Ports",
            "screenshots": "Screenshots",
        }
        return labels.get(test_name, test_name.replace("_", " ").title())

    def _add_findings_summary(self):
        """Add a summary table of findings across all domains."""
        self.document.add_heading("Key Findings", level=2)

        all_recs = []
        for domain, results in self.results.items():
            for result in results:
                for rec in result.recommendations:
                    all_recs.append({
                        "domain": domain,
                        "test": result.test_name,
                        "severity": rec["severity"],
                        "text": rec["text"],
                    })

        # Count by severity
        critical_count = sum(1 for r in all_recs if r["severity"] == "critical")
        warning_count = sum(1 for r in all_recs if r["severity"] == "warning")
        info_count = sum(1 for r in all_recs if r["severity"] == "info")

        # Findings count table
        headers = ["Severity", "Count"]
        rows = [
            ["Critical", str(critical_count)],
            ["Warning", str(warning_count)],
            ["Informational", str(info_count)],
        ]
        create_professional_table(self.document, headers, rows, col_widths=[3.0, 2.0])

        # Add top recommendations
        if all_recs:
            self.document.add_paragraph()
            self.document.add_heading("Top Recommendations", level=3)

            # Show critical and warning items first
            priority_recs = [r for r in all_recs if r["severity"] in ("critical", "warning")]
            for rec in priority_recs[:5]:
                self._add_recommendation_item(rec["severity"], rec["text"], rec["domain"])

    def _add_recommendation_item(self, severity: str, text: str, domain: str = ""):
        """Add a styled recommendation item."""
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)

        # Severity label
        label = SEVERITY_LABELS.get(severity, "INFO")
        color = SEVERITY_COLORS.get(severity, COLOR_INFO)

        run = p.add_run(f"[{label}] ")
        run.bold = True
        run.font.color.rgb = color
        run.font.size = Pt(9)
        run.font.name = FONT_HEADING

        # Domain
        if domain:
            run = p.add_run(f"({domain}) ")
            run.font.color.rgb = COLOR_MUTED
            run.font.size = Pt(9)
            run.italic = True

        # Text
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_BODY

    def _add_toc_placeholder(self):
        """Add a table of contents placeholder."""
        self.document.add_heading("Table of Contents", level=1)
        p = self.document.add_paragraph()
        run = p.add_run(
            "[Table of Contents - Please right-click and select "
            "'Update Field' to generate in Microsoft Word]"
        )
        run.font.color.rgb = COLOR_MUTED
        run.font.size = Pt(10)
        run.italic = True

        # Add a TOC field code
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)

        run = paragraph.add_run()
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run._r.append(instr_text)

        run = paragraph.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fld_char_end)

    def _add_domain_section(self, domain: str, results: list[TestResult]):
        """Add all test result sections for a domain."""
        self.document.add_heading(f"Results: {domain}", level=1)

        p = self.document.add_paragraph()
        run = p.add_run(
            f"The following sections detail the security assessment results for {domain}."
        )
        run.font.size = Pt(10)

        for result in results:
            self._add_test_result_section(result)

    def _add_test_result_section(self, result: TestResult):
        """Add a complete section for a single test result."""
        # Section heading
        test_titles = {
            "ssl_labs": "SSL/TLS Analysis",
            "mozilla_observatory": "Mozilla Observatory Analysis",
            "dns_lookup": "DNS Records & DNSSEC Validation",
            "email_auth": "Email Authentication (SPF, DKIM, DMARC)",
            "blacklist": "Domain & IP Blacklist Check",
            "subdomains": "Subdomain Enumeration",
            "security_headers": "Security Headers Analysis",
            "performance": "Performance Analysis",
            "cookies": "Cookie Security Analysis",
            "open_ports": "Open Ports Analysis",
            "screenshots": "Website Screenshot",
        }
        title = test_titles.get(result.test_name, result.test_name.replace("_", " ").title())
        self.document.add_heading(title, level=2)

        # Test description
        test_descriptions = {
            "ssl_labs": (
                "SSL Labs provides a comprehensive analysis of the SSL/TLS configuration "
                "of a web server. This includes certificate validation, protocol support, "
                "cipher suite analysis, and vulnerability checks."
            ),
            "mozilla_observatory": (
                "The Mozilla HTTP Observatory assesses web server security by analyzing "
                "HTTP response headers, redirect behavior, and security best practices. "
                "It provides a score and grade based on industry standards."
            ),
            "dns_lookup": (
                "DNS record analysis examines the domain's DNS configuration including "
                "A, AAAA, MX, NS, CNAME, and TXT records. DNSSEC validation checks "
                "whether the domain is protected against DNS spoofing attacks."
            ),
            "email_auth": (
                "Email authentication analysis checks for SPF (Sender Policy Framework), "
                "DKIM (DomainKeys Identified Mail), and DMARC (Domain-based Message "
                "Authentication) records. These protocols protect against email spoofing "
                "and phishing attacks."
            ),
            "blacklist": (
                "Blacklist checking verifies the domain and its IP address against "
                "major DNS-based blacklists (DNSBLs). Being listed can affect email "
                "deliverability and domain reputation."
            ),
            "subdomains": (
                "Subdomain enumeration discovers publicly visible subdomains using "
                "Certificate Transparency logs and DNS brute forcing. A large number "
                "of subdomains increases the attack surface."
            ),
            "security_headers": (
                "Security headers analysis examines HTTP response headers that provide "
                "additional security controls. Properly configured headers can prevent "
                "XSS, clickjacking, and other client-side attacks."
            ),
            "performance": (
                "Performance analysis evaluates page speed and best-practice web metrics "
                "using Lighthouse or the PageSpeed Insights API."
            ),
            "cookies": (
                "Cookie analysis checks for Secure, HttpOnly, and SameSite attributes "
                "to reduce session theft and cross-site request risks."
            ),
            "open_ports": (
                "Open ports analysis identifies exposed network services using nmap "
                "top-port scanning to reduce unnecessary internet exposure."
            ),
            "screenshots": (
                "Screenshot capture records the rendered web page for visual review and "
                "evidence in the assessment workflow."
            ),
        }
        desc = test_descriptions.get(result.test_name, "")
        if desc:
            p = self.document.add_paragraph()
            run = p.add_run(desc)
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_MUTED
            run.italic = True

        if not result.success:
            # Error state
            p = self.document.add_paragraph()
            run = p.add_run(f"Test Error: {result.error}")
            run.font.color.rgb = COLOR_CRITICAL
            run.bold = True
            return

        # Summary
        if result.summary:
            p = self.document.add_paragraph()
            run = p.add_run(result.summary)
            run.font.size = Pt(10)

        # Visuals
        self._add_visuals(result.visuals)

        # Tables
        for table_data in result.tables:
            self._add_result_table(table_data)

        # Recommendations
        if result.recommendations:
            self.document.add_heading("Recommendations", level=3)
            for rec in result.recommendations:
                self._add_recommendation_item(rec["severity"], rec["text"])

    def _add_visuals(self, visuals: dict[str, io.BytesIO]):
        """Add visual elements (charts, badges, gauges) to the document."""
        if not visuals:
            return

        # Grade badge and score gauge side-by-side if both exist
        badge_keys = [k for k in visuals if "badge" in k]
        gauge_keys = [k for k in visuals if "gauge" in k]
        chart_keys = [k for k in visuals if k not in badge_keys and k not in gauge_keys]

        # Add badges and gauges in a row
        badge_gauge_keys = badge_keys + gauge_keys
        if badge_gauge_keys:
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for key in badge_gauge_keys:
                buf = visuals[key]
                run = p.add_run()
                width = Inches(1.8) if "badge" in key else Inches(2.5)
                self._safe_add_picture(run, buf, width=width, warning_context=f"visual {key}")
                run.add_text("   ")  # spacing

        # Add charts
        for key in chart_keys:
            buf = visuals[key]
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            self._safe_add_picture(run, buf, width=Inches(5.0), warning_context=f"visual {key}")

    def _safe_add_picture(
        self,
        run,
        image_source: str | io.BytesIO,
        *,
        width=None,
        height=None,
        warning_context: str,
    ) -> bool:
        """Safely add an inline image and avoid breaking report generation."""
        if isinstance(image_source, str) and not os.path.exists(image_source):
            logger.warning("Skipping missing image (%s): %s", warning_context, image_source)
            return False

        try:
            if hasattr(image_source, "seek"):
                image_source.seek(0)
            run.add_picture(image_source, width=width, height=height)
            return True
        except Exception as exc:
            logger.warning("Failed to insert image (%s): %s", warning_context, exc)
            return False

    def _add_result_table(self, table_data: dict):
        """Add a table from test results."""
        title = table_data.get("title", "")
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        status_col = table_data.get("status_col")

        if not headers or not rows:
            return

        if title:
            self.document.add_heading(title, level=3)

        if status_col is not None:
            create_status_table(self.document, headers, rows, status_col=status_col)
        else:
            create_professional_table(self.document, headers, rows)

        # Add spacing after table
        self.document.add_paragraph()

    def _add_appendix(self):
        """Add the appendix section."""
        self._add_page_break()
        self.document.add_heading("Appendix", level=1)

        # Methodology
        self.document.add_heading("Methodology", level=2)
        p = self.document.add_paragraph()
        run = p.add_run(
            "This External Attack Surface Test (EAST) was conducted using automated "
            "scanning tools and publicly accessible APIs. The assessment focuses on "
            "externally visible security indicators and does not include internal "
            "network testing or application-level vulnerability scanning."
        )
        run.font.size = Pt(10)

        # Tools used
        self.document.add_heading("Tools and Services Used", level=2)
        tools = [
            ["SSL Labs API", "Qualys SSL Labs", "SSL/TLS configuration analysis"],
            ["Mozilla Observatory", "Mozilla Foundation", "HTTP security headers assessment"],
            ["dnspython", "DNS Toolkit", "DNS record lookups and DNSSEC validation"],
            ["DNS TXT Queries", "Direct DNS", "SPF, DKIM, and DMARC record analysis"],
            ["DNSBL Queries", "Multiple Providers", "IP and domain blacklist checking"],
            ["crt.sh (CT Logs)", "Sectigo", "Subdomain enumeration via Certificate Transparency"],
            ["HTTP Header Analysis", "Direct HTTP", "Security headers inspection"],
        ]
        create_professional_table(
            self.document,
            ["Tool", "Provider", "Purpose"],
            tools,
        )

        self.document.add_paragraph()

        # Scoring explanation
        self.document.add_heading("Scoring Guide", level=2)
        scoring = [
            ["A+ / A / A-", "Excellent", "Configuration meets or exceeds best practices"],
            ["B+ / B / B-", "Good", "Minor improvements recommended"],
            ["C+ / C / C-", "Fair", "Several improvements needed"],
            ["D / E", "Poor", "Significant security issues present"],
            ["F", "Fail", "Critical security deficiencies"],
        ]
        create_professional_table(
            self.document,
            ["Grade", "Rating", "Description"],
            scoring,
        )

        # Disclaimer
        self.document.add_paragraph()
        self.document.add_heading("Disclaimer", level=2)
        p = self.document.add_paragraph()
        run = p.add_run(
            "This report is provided for informational purposes only. The assessment "
            "was conducted at a point in time and results may change as configurations "
            "are updated. The findings should be used as guidance for improving security "
            "posture and should not be considered an exhaustive security audit. "
            "No destructive testing was performed during this assessment."
        )
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_MUTED

    def _add_page_break(self):
        """Add a page break to the document."""
        self.document.add_page_break()

    def _add_horizontal_line(self):
        """Add a horizontal line to the document."""
        p = self.document.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="dee2e6"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
