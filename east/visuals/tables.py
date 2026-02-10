"""Table formatting utilities for Word document generation."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Color scheme
HEADER_BG = "2c3e50"
HEADER_TEXT = "FFFFFF"
ROW_ALT_BG = "f8f9fa"
ROW_NORMAL_BG = "FFFFFF"
BORDER_COLOR = "dee2e6"
SUCCESS_COLOR = "28a745"
WARNING_COLOR = "ffc107"
CRITICAL_COLOR = "dc3545"
INFO_COLOR = "17a2b8"


def set_cell_shading(cell, color: str):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell.

    kwargs: top, bottom, left, right - each a dict with val, sz, color
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')

    for edge, props in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:val="{props.get("val", "single")}" '
            f'w:sz="{props.get("sz", "4")}" '
            f'w:space="0" '
            f'w:color="{props.get("color", BORDER_COLOR)}"/>'
        )
        tcBorders.append(element)

    tcPr.append(tcBorders)


def create_professional_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    """Create a professionally styled table in the document.

    Args:
        document: The python-docx Document object
        headers: List of header strings
        rows: List of row data (each row is a list of strings)
        col_widths: Optional list of column widths in inches
    """
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    # Style header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(HEADER_TEXT)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        set_cell_shading(cell, HEADER_BG)

    # Style data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        bg_color = ROW_ALT_BG if row_idx % 2 == 0 else ROW_NORMAL_BG

        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            set_cell_shading(cell, bg_color)

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def create_status_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    status_col: int = -1,
    col_widths: list[float] | None = None,
) -> None:
    """Create a table with color-coded status column.

    Args:
        document: The python-docx Document object
        headers: List of header strings
        rows: List of row data
        status_col: Index of the column to color-code (default: last column)
        col_widths: Optional column widths in inches
    """
    if status_col < 0:
        status_col = len(headers) + status_col

    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(HEADER_TEXT)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    status_map = {
        "pass": (SUCCESS_COLOR, "✓ Pass"),
        "passed": (SUCCESS_COLOR, "✓ Pass"),
        "present": (SUCCESS_COLOR, "✓ Present"),
        "supported": (SUCCESS_COLOR, "✓ Supported"),
        "valid": (SUCCESS_COLOR, "✓ Valid"),
        "secure": (SUCCESS_COLOR, "✓ Secure"),
        "not listed": (SUCCESS_COLOR, "✓ Not Listed"),
        "clean": (SUCCESS_COLOR, "✓ Clean"),
        "warn": (WARNING_COLOR, "⚠ Warning"),
        "warning": (WARNING_COLOR, "⚠ Warning"),
        "partial": (WARNING_COLOR, "⚠ Partial"),
        "fail": (CRITICAL_COLOR, "✗ Fail"),
        "failed": (CRITICAL_COLOR, "✗ Fail"),
        "missing": (CRITICAL_COLOR, "✗ Missing"),
        "not supported": (CRITICAL_COLOR, "✗ Not Supported"),
        "invalid": (CRITICAL_COLOR, "✗ Invalid"),
        "insecure": (CRITICAL_COLOR, "✗ Insecure"),
        "listed": (CRITICAL_COLOR, "✗ Listed"),
        "expired": (CRITICAL_COLOR, "✗ Expired"),
        "info": (INFO_COLOR, "ℹ Info"),
    }

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        bg_color = ROW_ALT_BG if row_idx % 2 == 0 else ROW_NORMAL_BG

        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "Calibri"

            if col_idx == status_col:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                status_key = str(cell_text).lower().strip()
                if status_key in status_map:
                    color_hex, display_text = status_map[status_key]
                    run.text = display_text
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(color_hex)
                else:
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            set_cell_shading(cell, bg_color)

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
