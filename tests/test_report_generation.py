import io
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE

from east.config import EASTConfig
from east.report import EASTReportGenerator
from east.tests.base import TestResult


class ReportGenerationTests(unittest.TestCase):
    def _result(self, domain: str, test_name: str, score: int, grade: str, success: bool = True) -> TestResult:
        return TestResult(
            test_name=test_name,
            domain=domain,
            success=success,
            score=score if success else None,
            grade=grade if success else "",
            summary="ok" if success else "failed",
            error="boom" if not success else "",
            visuals={"grade_badge": io.BytesIO(b"not-an-image")} if not success else {},
        )

    def test_assessment_overview_grid_renders_for_many_tests(self):
        config = EASTConfig.default()
        config.domains = ["example.com", "example.org", "example.net"]
        config.branding.logo = "assets/missing-logo.png"

        generator = EASTReportGenerator(config)

        test_names = [
            "ssl_labs", "mozilla_observatory", "dns_lookup", "email_auth", "blacklist",
            "subdomains", "security_headers", "performance", "cookies", "open_ports",
        ]

        for domain in config.domains:
            results = [self._result(domain, name, 85, "B") for name in test_names]
            results.append(self._result(domain, "screenshots", 0, "", success=False))
            generator.add_results(domain, results)

        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "report.docx"
            generator.generate(str(output))
            self.assertTrue(output.exists())

            doc = Document(str(output))
            joined_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Assessment Overview", joined_text)
            self.assertGreaterEqual(len(doc.tables), 4)  # overview tables + other report tables

            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
            self.assertIn("Screenshots (ERROR)", table_text)

            overview_table = doc.tables[0]
            for row in overview_table.rows:
                self.assertEqual(row.height_rule, WD_ROW_HEIGHT_RULE.EXACTLY)
                self.assertAlmostEqual(row.height.inches, 0.7, places=2)


if __name__ == "__main__":
    unittest.main()
